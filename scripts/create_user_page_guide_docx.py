from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
source = (ROOT / "docs" / "USER_PAGE_GUIDE.md").read_text()
document = Document()
section = document.sections[0]
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

for line in source.splitlines():
    if not line.strip():
        continue
    if line.startswith("# "):
        paragraph = document.add_heading(line[2:], level=0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith("## "):
        document.add_heading(line[3:], level=1)
    elif line.startswith("### "):
        document.add_heading(line[4:], level=2)
    elif line.startswith("- "):
        paragraph = document.add_paragraph(style="List Bullet")
        for index, part in enumerate(line[2:].split("**")):
            run = paragraph.add_run(part)
            run.bold = index % 2 == 1
    elif line.startswith("> "):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.35)
        run = paragraph.add_run(line[2:])
        run.italic = True
    else:
        document.add_paragraph(line)

style = document.styles["Normal"]
style.font.name = "Aptos"
style.font.size = Pt(10)
document.save(ROOT / "docs" / "Bazaar_User_Page_Guide.docx")
